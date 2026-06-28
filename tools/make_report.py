# -*- coding: utf-8 -*-
# مولّد تقارير Office احترافي عبر openpyxl (Excel) و python-docx (Word).
# يقرأ JSON من ملف (argv[1]) أو stdin: {format, out, title, period, school:{name,address}, headers:[], widths:[], landscape, rows:[{type, cells:[]}]}
import sys, os, json, glob

# تأكّد أن مجلد مكتبات المستخدم (حيث openpyxl/python-docx) على المسار — مهم عند تشغيل
# Apache للسكربت بمستخدم مختلف لا يرى site-packages الخاص بالمستخدم.
for _base in [r'C:\Users\user\AppData\Roaming\Python', os.path.expandvars(r'%APPDATA%\Python')]:
    for _sp in glob.glob(os.path.join(_base, 'Python*', 'site-packages')):
        if _sp not in sys.path:
            sys.path.insert(0, _sp)


def is_num(v):
    if isinstance(v, bool):
        return False
    if isinstance(v, (int, float)):
        return True
    if isinstance(v, str):
        s = v.replace(',', '')
        try:
            float(s); return s != ''
        except ValueError:
            return False
    return False


def numv(v):
    return float(str(v).replace(',', ''))


def fmt_num(v):
    n = numv(v)
    return f'{int(n):,}' if n == int(n) else (f'{n:,.2f}'.rstrip('0').rstrip('.'))


def make_xlsx(spec, out):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    from openpyxl.utils import get_column_letter
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.sheet_view.rightToLeft = True
    ws.title = (spec.get('title', 'تقرير') or 'تقرير')[:28]
    rows = spec.get('rows', [])
    headers = spec.get('headers', [])
    cols = max([len(headers)] + [len(r['cells']) for r in rows] + [1])
    thin = Side(style='thin', color='94A3B8')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    head_fill = PatternFill('solid', fgColor='1E3A8A')
    tot_fill = PatternFill('solid', fgColor='F3F4F6')
    sec_fill = PatternFill('solid', fgColor='E0E7FF')
    r = 1

    def merge_row(text, font):
        nonlocal r
        c = ws.cell(r, 1, text)
        c.font = font
        c.alignment = Alignment(horizontal='center', vertical='center')
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=cols)
        r += 1

    merge_row(spec.get('title', ''), Font(bold=True, size=15, color='1E3A8A'))
    if spec.get('period'):
        merge_row(spec['period'], Font(bold=True, size=12))
    sch = spec.get('school') or {}
    if sch.get('name'):
        extra = (' — ' + sch['address']) if sch.get('address') else ''
        merge_row(sch['name'] + extra, Font(bold=True, size=12))
    r += 1
    header_row = r
    for j, h in enumerate(headers, 1):
        c = ws.cell(header_row, j, h)
        c.font = Font(bold=True, color='FFFFFF')
        c.fill = head_fill
        c.border = border
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    if headers:
        r += 1
    for row in rows:
        cells = row['cells']
        t = row['type']
        if t == 'section':
            c = ws.cell(r, 1, cells[0] if cells else '')
            c.font = Font(bold=True)
            c.fill = sec_fill
            c.alignment = Alignment(horizontal='right', vertical='center')
            c.border = border
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=cols)
            for j in range(2, cols + 1):
                cc = ws.cell(r, j)
                cc.border = border
                cc.fill = sec_fill
            r += 1
            continue
        bold = (t == 'total')
        fill = tot_fill if t == 'total' else None
        for j in range(1, cols + 1):
            v = cells[j - 1] if j - 1 < len(cells) else ''
            cell = ws.cell(r, j)
            if is_num(v) and v != '':
                nv = numv(v)
                cell.value = nv
                cell.number_format = '#,##0' if nv == int(nv) else '#,##0.0'
                cell.alignment = Alignment(horizontal='left', vertical='center')
            else:
                cell.value = v
                cell.alignment = Alignment(horizontal='right', vertical='center')
            cell.border = border
            if bold:
                cell.font = Font(bold=True)
            if fill:
                cell.fill = fill
        r += 1
    widths = spec.get('widths') or []
    for j in range(1, cols + 1):
        w = widths[j - 1] if j - 1 < len(widths) else (28 if j == 2 else 15)
        ws.column_dimensions[get_column_letter(j)].width = w
    ws.freeze_panes = ws.cell(header_row + 1, 1)
    ws.page_setup.orientation = 'landscape' if spec.get('landscape') else 'portrait'
    ws.page_setup.paperSize = 9
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_options.gridLines = False
    wb.save(out)


def make_docx(spec, out):
    import docx
    from docx.shared import Pt, RGBColor, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    d = docx.Document()
    sec = d.sections[0]
    if spec.get('landscape'):
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = Cm(29.7), Cm(21)
    else:
        sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(1)
    land = spec.get('landscape')

    def heading(text, size, color=None):
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        run.font.rtl = True

    heading(spec.get('title', ''), 15, '1E3A8A')
    if spec.get('period'):
        heading(spec['period'], 12)
    sch = spec.get('school') or {}
    if sch.get('name'):
        extra = (' — ' + sch['address']) if sch.get('address') else ''
        heading(sch['name'] + extra, 12)
    d.add_paragraph()
    rows = spec.get('rows', [])
    headers = spec.get('headers', [])
    cols = max([len(headers)] + [len(r['cells']) for r in rows] + [1])

    # عرض الأعمدة المتناسب: يملأ عرض الصفحة بالكامل بلا تكسير الأرقام على سطرين.
    content_cm = 27.7 if land else 19.0
    raw = spec.get('widths') or []
    w = [(raw[i] if i < len(raw) else (28 if i == 1 else 15)) for i in range(cols)]
    tot_w = sum(w) or cols
    col_tw = [max(360, int(content_cm / 2.54 * 1440 * wi / tot_w)) for wi in w]  # twips لكل عمود
    total_tw = sum(col_tw)
    # خط متأقلم: الجداول العريضة (كثيرة الأعمدة) تحتاج خطاً أصغر ليسع الأرقام
    base_sz = 8 if (land and cols >= 14) else (9 if land else 11)

    table = d.add_table(rows=0, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    tblPr = table._tbl.tblPr
    tblPr.append(OxmlElement('w:bidiVisual'))
    # تخطيط ثابت + عرض كلّي محدّد (يمنع python-docx من إعادة التوزيع بالتساوي)
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(total_tw)); tblW.set(qn('w:type'), 'dxa'); tblPr.append(tblW)
    # شبكة الأعمدة بعرضها المحدّد
    grid = table._tbl.find(qn('w:tblGrid'))
    if grid is None:
        grid = OxmlElement('w:tblGrid'); table._tbl.insert(list(table._tbl).index(tblPr) + 1, grid)
    for gc in list(grid):
        grid.remove(gc)
    for tw in col_tw:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(tw)); grid.append(gc)

    def shade(cell, color):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement('w:shd')
        sh.set(qn('w:val'), 'clear')
        sh.set(qn('w:fill'), color)
        tcPr.append(sh)

    def setcell(cell, text, bold=False, white=False, align='right', fill=None, numeric=False, sz=None):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.RIGHT
        if not numeric:
            p._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(sz if sz else base_sz)
        # الأرقام تبقى لاتينية (غربية) ولا تُكسَّر: لا تعليم RTL على الأرقام
        run.font.rtl = (not numeric)
        if white:
            run.font.color.rgb = RGBColor.from_string('FFFFFF')
        if fill:
            shade(cell, fill)

    section_rows = set()
    ri = -1
    if headers:
        rc = table.add_row().cells
        ri += 1
        for j, h in enumerate(headers):
            setcell(rc[j], h, bold=True, white=True, align='center', fill='1E3A8A', sz=base_sz)
    for row in rows:
        cells = row['cells']
        t = row['type']
        rc = table.add_row().cells
        ri += 1
        if t == 'section':
            section_rows.add(ri)
            rc[0].merge(rc[cols - 1])
            setcell(rc[0], cells[0] if cells else '', bold=True, align='right', fill='E0E7FF')
            continue
        bold = (t == 'total')
        fill = 'F3F4F6' if t == 'total' else None
        for j in range(cols):
            v = cells[j] if j < len(cells) else ''
            if is_num(v) and v != '':
                setcell(rc[j], fmt_num(v), bold=bold, align='center', fill=fill, numeric=True)
            else:
                setcell(rc[j], v, bold=bold, align='right', fill=fill)

    # ثبّت عرض كل خلية (عدا صفوف الأقسام المدموجة) ليحترم Word الأعمدة
    for idx, trow in enumerate(table.rows):
        if idx in section_rows:
            continue
        for j in range(cols):
            try:
                trow.cells[j].width = Twips(col_tw[j])
            except Exception:
                pass
    d.save(out)


def main():
    # JSON من ملف (argv[1]) إن وُجد، وإلا من stdin
    if len(sys.argv) > 1:
        with open(sys.argv[1], 'r', encoding='utf-8') as f:
            spec = json.load(f)
    else:
        spec = json.load(sys.stdin)
    if spec.get('format') == 'docx':
        make_docx(spec, spec['out'])
    else:
        make_xlsx(spec, spec['out'])
    sys.stdout.write('OK')


if __name__ == '__main__':
    main()
